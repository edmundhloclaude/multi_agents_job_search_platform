"""Tests: document ingestion, interactive StrategySession, and the web endpoints."""

import base64
import io
import json
import os
import sys
import threading
import urllib.request

import pytest

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from jobsearch.agents.doc_ingest import extract_text
from jobsearch.agents.strategy import load_criteria_from_strategy
from jobsearch.agents.strategy_session import StrategySession
from jobsearch.llm.base import LLM
from jobsearch.models import Accomplishment, AccomplishmentBank
from jobsearch.strategy_web import create_server


class FakeLLM(LLM):
    model = "fake"
    def __init__(self, criteria, reply="Got it.", entries=None):
        self.criteria, self.reply, self.calls = criteria, reply, 0
        self.entries = entries or []
    def complete_text(self, *a, **k):
        return ""
    def complete_json(self, system, user, *, temperature=0.0):
        self.calls += 1
        if "entries" in system.lower():   # llm_draft_bank_entries prompt
            return {"entries": self.entries}
        return {"reply": self.reply, "criteria": self.criteria}


@pytest.fixture
def bank():
    return AccomplishmentBank(
        name="Jane Doe", contact={},
        accomplishments=[Accomplishment("Acme", "Senior Engineer", "2019", "2023",
                         "Built distributed systems in Python.", metrics=[],
                         skills=["python", "distributed systems"])],
        skills=["python", "distributed systems"], credentials=[])


# --------------------------------------------------------------------------- #
# Document ingestion
# --------------------------------------------------------------------------- #
def test_ingest_text_and_markdown():
    assert "hello" in extract_text("a.txt", b"hello world")
    assert "# Title" in extract_text("a.md", b"# Title\nbody")


def test_ingest_docx_roundtrip():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Staff Engineer at Acme")
    doc.add_paragraph("Python, Kubernetes")
    buf = io.BytesIO(); doc.save(buf)
    text = extract_text("resume.docx", buf.getvalue())
    assert "Staff Engineer at Acme" in text and "Kubernetes" in text


# --------------------------------------------------------------------------- #
# StrategySession
# --------------------------------------------------------------------------- #
def test_session_send_updates_criteria(bank):
    llm = FakeLLM({"target_roles": ["Staff Engineer"], "seniority": ["staff"],
                   "must_haves": ["python"], "keywords_boost": ["kubernetes"],
                   "comp_min": 200000, "remote_ok": True})
    s = StrategySession(llm, profile={}, bank=bank)
    st = s.send("I want staff backend roles")
    assert st["criteria"]["target_roles"] == ["Staff Engineer"]
    assert st["criteria"]["comp_min"] == 200000
    # both the user turn and assistant reply are recorded
    roles = [m["role"] for m in s.messages]
    assert roles == ["user", "assistant"]
    assert st["last_reply"] == "Got it."


def test_session_flags_ungrounded_must_have(bank):
    llm = FakeLLM({"must_haves": ["python", "rust"]})  # rust not in bank
    s = StrategySession(llm, profile={}, bank=bank)
    st = s.send("go")
    assert "rust" in st["ungrounded_must_haves"]
    assert "python" not in st["ungrounded_must_haves"]


def test_session_document_triggers_turn(bank):
    llm = FakeLLM({"target_roles": ["ML Engineer"]})
    s = StrategySession(llm, profile={}, bank=bank)
    s.add_document("resume.txt", "Machine learning engineer, TensorFlow")
    st = s.note_document("resume.txt")
    assert "resume.txt" in st["documents"]
    assert st["criteria"]["target_roles"] == ["ML Engineer"]
    assert llm.calls == 1


def test_draft_and_commit_bank_entries(tmp_path, bank):
    import yaml
    bank_path = tmp_path / "accomplishment_bank.yaml"
    # start from an existing bank with identity we must not clobber
    bank_path.write_text(yaml.safe_dump({
        "name": "Jane Doe", "contact": {"email": "j@e.com"},
        "accomplishments": [], "skills": ["python"]}))
    entries = [{"employer": "Stripe", "title": "Senior SWE", "start_date": "2020",
                "end_date": "2024", "text": "Built payments infra handling 50k req/s.",
                "metrics": ["50k"], "skills": ["go", "kubernetes"]}]
    llm = FakeLLM({"target_roles": ["SWE"]}, entries=entries)
    s = StrategySession(llm, profile={}, bank=bank, bank_path=str(bank_path))
    s.add_document("resume.txt", "Stripe Senior SWE ... 50k req/s ... Go, Kubernetes")

    drafted = s.draft_bank_entries()
    assert drafted and drafted[0]["employer"] == "Stripe"

    added = s.commit_bank_entries(drafted)
    assert added == 1
    written = yaml.safe_load(bank_path.read_text())
    assert written["name"] == "Jane Doe"                       # identity preserved
    assert any(a["employer"] == "Stripe" for a in written["accomplishments"])
    assert "go" in written["skills"] and "kubernetes" in written["skills"]
    # in-memory bank refreshed so grounding sees the new skills
    assert "kubernetes" in s.bank.known_skills()


def test_commit_without_bank_path_is_noop(bank):
    s = StrategySession(FakeLLM({}), profile={}, bank=bank, bank_path=None)
    assert s.commit_bank_entries([{"employer": "X", "text": "y"}]) == 0


def test_session_save_roundtrips(tmp_path, bank):
    llm = FakeLLM({"target_roles": ["Staff Engineer"], "must_haves": ["python"],
                   "comp_min": 190000})
    s = StrategySession(llm, profile={}, bank=bank)
    s.send("staff roles please")
    path = tmp_path / "strategy.md"
    s.save(path)
    parsed = load_criteria_from_strategy(path)
    assert "python" in parsed.must_haves
    assert parsed.comp_min == 190000
    assert "interactive session" in path.read_text()


# --------------------------------------------------------------------------- #
# Web endpoints (real HTTP on an ephemeral port, injected fake session)
# --------------------------------------------------------------------------- #
@pytest.fixture
def server(tmp_path, bank):
    llm = FakeLLM({"target_roles": ["Staff Engineer"], "must_haves": ["python"]})
    session = StrategySession(llm, profile={}, bank=bank)
    srv = create_server(session, str(tmp_path / "strategy.md"), host="127.0.0.1", port=0)
    port = srv.server_address[1]
    t = threading.Thread(target=srv.serve_forever, daemon=True); t.start()
    yield f"http://127.0.0.1:{port}", srv, tmp_path
    srv.shutdown(); srv.server_close()


def _post(base, path, body):
    req = urllib.request.Request(base + path, method="POST",
                                 data=json.dumps(body).encode(),
                                 headers={"Content-Type": "application/json"})
    return json.loads(urllib.request.urlopen(req).read())


def test_web_serves_ui_and_chat(server):
    base, _, _ = server
    html = urllib.request.urlopen(base + "/").read().decode()
    assert "Strategy Advisor" in html
    st = _post(base, "/api/message", {"text": "hi"})
    assert st["criteria"]["target_roles"] == ["Staff Engineer"]
    assert any(m["role"] == "assistant" for m in st["messages"])


def test_web_upload_and_save(server):
    base, srv, tmp_path = server
    b64 = base64.b64encode(b"resume text: python, distributed systems").decode()
    st = _post(base, "/api/upload", {"name": "resume.txt", "b64": b64})
    assert "resume.txt" in st["documents"]
    r = _post(base, "/api/save", {})
    assert os.path.exists(r["saved"])
    assert "python" in (tmp_path / "strategy.md").read_text()
